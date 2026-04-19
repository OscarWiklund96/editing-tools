"""Extract text content from DOCX files using python-docx."""

import re

from docx import Document
from docx.oxml.ns import qn


def extract_text(filepath: str) -> str:
    """Extract and return all text from the given DOCX file."""
    try:
        doc = Document(filepath)
    except Exception as e:
        raise ValueError(f"Could not open DOCX file '{filepath}': {e}") from e

    parts: list[str] = []

    # Body paragraphs
    for para in doc.paragraphs:
        text = para.text
        if text.strip():
            parts.append(text)
        else:
            # Preserve intentional blank-line separators as empty entries
            # so adjacent non-empty paragraphs get double-newline separation.
            parts.append("")

    # Tables
    for table in doc.tables:
        table_rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_rows.append("\t".join(cells))
        parts.append("\n".join(table_rows))

    # Join: consecutive non-empty blocks get "\n\n"; empty entries act as
    # separators; consecutive empties are collapsed.
    lines: list[str] = []
    prev_empty = False
    for part in parts:
        if part == "":
            if not prev_empty:
                lines.append("")
            prev_empty = True
        else:
            lines.append(part)
            prev_empty = False

    raw = "\n\n".join(
        block for block in "\n".join(lines).split("\n\n") if block.strip()
    )
    return _normalise_whitespace(raw)


def _normalise_whitespace(text: str) -> str:
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = "\n".join(line.rstrip() for line in text.splitlines())
    return text.strip()
